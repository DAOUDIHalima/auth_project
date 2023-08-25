# stockManagement/views.py
from django.shortcuts import render, get_object_or_404, redirect
from .models import Medicine,MedicineStockMovement
from .form import MedicineForm, PatientForm
from django.core.paginator import Paginator
from django.db.models import Q
from django.http import JsonResponse
from datetime import date
from django.shortcuts import render
from django.db.models import Q
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.core.paginator import Paginator
from django.shortcuts import render, redirect, get_object_or_404
from django.templatetags.static import static 
from .form import PatientForm
from .models import  DossierMedical, Patient,User
from datetime import timedelta
from .form import NotesForm, PatientForm, InfoForm, SystematiqueForm, VisiteRepriseForm, SurveillanceForm, \
    LaboratoireForm
from .models import Notes, Patient, User, Analyse, systematique, info, surveillance, Laboratoire
from django.views.generic import DetailView
import locale
from datetime import date
from django.urls import reverse
from django.http import HttpResponse, JsonResponse, FileResponse


# Create your views here.
from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect
def logout(request):
    return render(request, 'registration/login.html')

def login_view(request):
    msg = ''
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(username=username, password=password)
        print(user)
        if user is not None:
            print('hello !')
            login(request, user)
            if user.is_medecin:
                return redirect('medecin')
            elif user.is_infirmier:
                return redirect('infirmier')
            else:
                return redirect('admin')  # Redirect to appropriate admin page
        else:
            msg = 'Username or password incorrect'
    return render(request, 'registration/login.html', {'msg': msg})

def dashboard_inf(request):
    patient_count = Patient.objects.count()
    medicine_count = Medicine.objects.count()
    outmedicine_count = OutStockMedicine.objects.count()

    ambulanceItem_count =AmbulanceItem.objects.count()
    ambulanceMedicine_count = AmbulanceMedicine.objects.count()
    visit_count = VisiteMedicaleSpontanee.objects.count()

    infirmier_count = User.objects.filter(is_infirmier=True).count()
    today = date.today()
    expired_medicines = Medicine.objects.filter(expiration_date__lt=today)
    soon_expiring_medicines = Medicine.objects.filter(
        expiration_date__gt=today,
        expiration_date__lte=today + timedelta(days=60)
    )
    context = {
        'patient_count': patient_count,
        'medicine_count': medicine_count,
        'outmedicine_count': outmedicine_count,
        'ambulanceItem_count': ambulanceItem_count,
        'infirmier_count': infirmier_count,
        'visit_count': visit_count,
        'ambulanceMedicine_count': ambulanceMedicine_count,
        'expired_medicines': expired_medicines,  # Pass the list of expired medicines to the context
        'soon_expiring_medicines': soon_expiring_medicines,
    }
    
    return render(request,'infirmier/dashboard.html',context)


def dashboard_med(request):
    patient_count = Patient.objects.count()
    medecin_count = User.objects.filter(is_medecin=True).count()

    # Pass the count to the template context
    context = {
        'patient_count': patient_count,
        'medecin_count': medecin_count, 
            }
    return render(request,'Medecin/dashboard.html',context)

def medecin(request):
    return render(request, 'Medecin/medecin.html')

def infirmier(request):
    return render(request, 'infirmier/infirmier.html')

def admin(request):
    return render(request, 'admin.html')

def dossier(request ):
    dossier = DossierMedical.objects.all()
    return render(request, "dossier1/dossier.html", {'dossier': dossier})

######Patient

def patient_list(request):
    user = request.user

    if user.is_medecin:
        user.template_name = 'Medecin/medecin.html'  # Set the template_name for medecin
    elif user.is_infirmier:
        user.template_name = 'Infirmier/infirmier.html'  # Set the template_name for infirmier
    patients = Patient.objects.all().order_by('name')
    return render(request, 'Patient/patients.html', {'patients': patients})
from django.core.files import File
import os

def create_patient(request):
    user = request.user

    if user.is_medecin:
        user.template_name = 'Medecin/medecin.html'  # Set the template_name for medecin
    elif user.is_infirmier:
        user.template_name = 'Infirmier/infirmier.html'  # Set the template_name for infirmier
    
    if request.method == 'POST':
        form = PatientForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            patient = form.save(commit=False)  # Create a patient instance without saving to the database

            if patient.area == "CCR":
                patient.risk.save("risk_file_CCR.pdf", open('data/risk_files/risk_file_CCR.pdf', 'rb'))
            elif patient.area == "SOLAR FIELD":
                patient.risk.save("risk_file_SolarField.pdf", open('data/risk_files/risk_file_SolarField.pdf', 'rb'))
            elif patient.area == "ADMINISTRATION OFFICE":
                patient.risk.save("risk_file_AdminOffice.pdf", open('data/risk_files/risk_file_AdminOffice.pdf', 'rb'))
            elif patient.area == "WTP":
                patient.risk.save("risk_file_WTP.pdf", open('data/risk_files/risk_file_WTP.pdf', 'rb'))
            elif patient.area == "WORKSHOP":
                patient.risk.save("risk_file_Workshop.pdf", open('data/risk_files/risk_file_Workshop.pdf', 'rb'))
            elif patient.area == "WAREHOUSE":
                patient.risk.save("risk_file_Warehouse.pdf", open('data/risk_files/risk_file_Warehouse.pdf', 'rb'))

            # Set any other attributes of the patient instance if needed
            patient.save()  # Now save the patient instance with the associated risk file

            return redirect('patients')
        else:
            print("Form errors:", form.errors)  # Debug print to check form errors
    else:
        form = PatientForm()

    return render(request, 'Patient/create_patient.html', {'form': form})

def update_patient(request,pk):
    user = request.user

    if user.is_medecin:
        user.template_name = 'Medecin/medecin.html'  # Set the template_name for medecin
    elif user.is_infirmier:
        user.template_name = 'Infirmier/infirmier.html'  # Set the template_name for infirmier
    patient=get_object_or_404(Patient,pk=pk)
    if request.method =='POST':
        form=PatientForm(request.POST or None, request.FILES or None ,instance=patient)
        if form.is_valid():
            form.save()
            return redirect('patients')
    else:
        form=PatientForm(instance=patient)
    return render(request,'Patient/update_patient.html',{'form':form,'patient':patient})

def view_patient(request, pk):
    user = request.user

    if user.is_medecin:
        user.template_name = 'Medecin/medecin.html'  # Set the template_name for medecin
    elif user.is_infirmier:
        user.template_name = 'Infirmier/infirmier.html'  # Set the template_name for infirmier
    patient = get_object_or_404(Patient, pk=pk)
    form = PatientForm(instance=patient)
    return render(request, 'Patient/view_patient.html', {'form': form , 'patient':patient})

def delete_patient(request, pk):
    user = request.user

    if user.is_medecin:
        user.template_name = 'Medecin/medecin.html'  # Set the template_name for medecin
    elif user.is_infirmier:
        user.template_name = 'Infirmier/infirmier.html'  # Set the template_name for infirmier
    patient = get_object_or_404(Patient,pk=pk)
    patient.delete()
    return redirect('patients')

def home(request):
    return render(request, 'home.html')
def index(request):
    return render(request, 'index.html')
def list_patients(request):
    contact_list = Patient.objects.all()
    paginator = Paginator(contact_list, 5)  # Show 25 contacts per page.

    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    return render(request, 'Patient/patients.html', {"page_obj": page_obj})



def search_Patient(request):
    if request.method == 'POST':
        query = request.POST.get('keyword')
        # Filter patients based on 'nom' or 'entreprise' containing the query
        patients = Patient.objects.filter(Q(name__icontains=query) | Q(company__icontains=query))
        return render(request, 'Patient/patients.html', {'patients': patients})
    else:
        # If the form is not submitted, show all patients
        patients = Patient.objects.all()
        return render(request, 'Patient/patients.html', {'patients': patients})
from django.shortcuts import render
from .models import MedicineStockMovement

def movement_list(request):
    movements = MedicineStockMovement.objects.all()
    return render(request, 'MovementList/movement_list.html', {'movements': movements})
    
from django.views.generic import DetailView
from .models import Patient

class PatientDetailView(DetailView):
    model = Patient
    template_name = 'view_patient.html'
    context_object_name = 'patient'
#########################################################################################    


def medicine_list(request):
    # Retrieve all medicine items from the database and pass them to the template
    medicines = Medicine.objects.all().order_by('expiration_date')
    return render(request, 'Medicine/IN/medicines.html', {'medicines': medicines})

def create_medicine(request):
    if request.method == 'POST':
        form = MedicineForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            form.save()
            return redirect('medicines')
        else:
            print("Form errors:", form.errors)  # Debug print to check form errors
    else:
        form = MedicineForm()

    return render(request, 'Medicine/IN/create_medicine.html', {'form': form})

def update_medicine(request, pk):
    medicine = get_object_or_404(Medicine, pk=pk)
    if request.method == 'POST':
        form = MedicineForm(request.POST or None, request.FILES or None, instance=medicine)
        if form.is_valid():
            form.save()
            return redirect('medicines')
    else:
        form = MedicineForm(instance=medicine)
    return render(request, 'Medicine/IN/update_medicine.html', {'form': form, 'medicine': medicine})

def view_medicine(request, pk):
    medicine = get_object_or_404(Medicine, pk=pk)
    form = MedicineForm(instance=medicine)
    return render(request, 'Medicine/IN/view_medicine.html', {'form': form, 'medicine': medicine})
def delete_medicine(request, pk):
    medicine_del = get_object_or_404(Medicine, pk=pk)
    medicine_del.delete()
    return redirect('medicines')

def home(request):
    return render(request, 'home.html')


def search_Medicine(request):
    if request.method == 'POST':
        query = request.POST.get('keyword')
        # Filter medicines based on 'name' or 'manufacturer' containing the query
        medicines = Medicine.objects.filter(Q(name__icontains=query) | Q(manufacturer__icontains=query))
        return render(request, 'Medicine/IN/medicines.html', {'medicines': medicines})
    else:
        # If the form is not submitted, show all medicines
        medicines = Medicine.objects.all()
        return render(request, 'Medicine/IN/medicines.html', {'medicines': medicines})

def expired_medicines_notification(request):
    # Get all expired medicines
    expired_medicines = Medicine.objects.filter(expiration_date__lt=date.today())
    return {'expired_medicines'}
##################################
import io
import csv
import openpyxl
from io import BytesIO
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx import Document
import requests

def download_medicines(request, format):
    medicines = Medicine.objects.all()

    if format == 'pdf':
        buffer = io.BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=letter)
        company_logo_url = request.build_absolute_uri(static('nomac.jpg'))  # Build the absolute URL
        logo = Image(company_logo_url, width=100, height=70)  
        title_style = getSampleStyleSheet()["Title"]
        title1 = Paragraph("Medical Inventory Report", title_style)
        logo_title_table = Table([[logo, title1]], colWidths=[5, '*'])
        # Add the logo, title, and current date within a spacer to center align them
        elements= [logo_title_table,Spacer(0, 10),]
        # Table data for PDF
        data = [['S.No','Name', 'Minimum', 'Actual', 'Expiry date', 'Remarks']]
        i=0
        for medicine in medicines:
            i+=1
            data.append([i, medicine.name, medicine.quantity, medicine.expiration_date, medicine.status])

        # Define table and table style
        column_widths = [25,150, 60, 60, 80,80]
        table = Table(data,colWidths=column_widths)
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, '#888888'),
        ]))

        # Build the PDF document
      
        elements.append(table)
        pdf.build(elements)

        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="medicine_list.pdf"'
        return response

    elif format == 'excel':
    # Excel Generation Logic
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = 'Medicine List'

    # Table data for Excel
        data = [['Name', 'Quantity',  'Expiration Date',  'Dosage', 'Reception Date', 'Status']]
        for medicine in medicines:
             data.append([medicine.name, medicine.quantity,  medicine.expiration_date.strftime('%Y-%m-%d'),medicine.dosage, medicine.reception_date.strftime('%Y-%m-%d'),medicine.status])

    # Write data to the sheet
        for row in data:
            sheet.append(row)

    # Save the virtual workbook
        virtual_workbook = BytesIO()
        workbook.save(virtual_workbook)
        virtual_workbook.seek(0)

        response = HttpResponse(virtual_workbook.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="medicine_list.xlsx"'
        return response

    elif format == 'doc':
        # Word Document Generation Logic
        response = HttpResponse(content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename="medicine_list.doc"'

        # Create a new Word document
        doc = Document()

        # Add logo and title
        company_logo_url = request.build_absolute_uri(static('nomac.jpg'))  # Build the absolute URL
        logo_stream = BytesIO(requests.get(company_logo_url).content)
        doc.add_picture(logo_stream, width=Inches(2), height=Inches(1))  # Adjust width and height as needed
        title = doc.add_paragraph("Medical Inventory Report")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a table with headers
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'Name'
        headers[1].text = 'Quantity'
        headers[2].text = 'Expiration Date'
        headers[3].text = 'Dosage'
        headers[4].text = 'Reception Date'
        headers[5].text = 'Status'

        # Add data rows to the table
        for medicine in medicines:
            row_cells = table.add_row().cells
            row_cells[0].text = medicine.name
            row_cells[1].text = str(medicine.quantity)
            row_cells[2].text = str(medicine.expiration_date)
            row_cells[3].text = medicine.dosage
            row_cells[4].text = str(medicine.reception_date)
            row_cells[5].text = str(medicine.status)

        # Save the Word document to the response
        doc.save(response)
        return response
    else:
        return HttpResponse("Invalid format.")
    
# ambulanceapp/views.py

from django.shortcuts import render, get_object_or_404, redirect
from .models import AmbulanceItem ,AmbulanceMedicine
from .form import AmbulanceItemForm ,AmbulanceMedicineForm
from django.core.paginator import Paginator

def ambulance_item_list(request):
    # Retrieve all medicine items from the database and pass them to the template
    ambulance_items = AmbulanceItem.objects.all().order_by('name')
    return render(request, 'Ambulance/AmbulanceItems/ambulance_Item.html', {'ambulance_items': ambulance_items})

def create_ambulance_item(request):
    if request.method == 'POST':
        form = AmbulanceItemForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            form.save()
            return redirect('ambulance_items')
        else:
            print("Form errors:", form.errors)  # Debug print to check form errors
    else:
        form = AmbulanceItemForm()

    return render(request, 'Ambulance/AmbulanceItems/create_ambulance_Item.html', {'form': form})

def update_ambulance_item(request, pk):
    ambulance_item = get_object_or_404(AmbulanceItem, pk=pk)
    if request.method == 'POST':
        form = AmbulanceItemForm(request.POST or None, request.FILES or None, instance=ambulance_item)
        if form.is_valid():
            form.save()
            return redirect('ambulance_items')
    else:
        form = AmbulanceItemForm(instance=ambulance_item)
    return render(request, 'Ambulance/AmbulanceItems/update_ambulance_Item.html', {'form': form, 'ambulance_item': ambulance_item})

def view_ambulance_item(request, pk):
    ambulance_item = get_object_or_404(AmbulanceItem, pk=pk)
    form = AmbulanceItemForm(instance=ambulance_item)
    return render(request, 'Ambulance/AmbulanceItems/view_ambulance_Item.html', {'form': form, 'ambulance_item': ambulance_item})

def delete_ambulance_item(request, pk):
    ambulance_item = get_object_or_404(AmbulanceItem, pk=pk)
    ambulance_item.delete()
    return redirect('ambulance_items')


def list_ambulance_items(request):
    # Retrieve all ambulance items from the database and pass them to the template
    ambulance_items = AmbulanceItem.objects.all().order_by('name')
    # Create a paginator with 10 items per page
    paginator = Paginator(ambulance_items, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    return render(request, 'list_ambulance_Item.html', {"page_obj": page_obj})

def search_AmbulanceItem(request):
    if request.method == 'POST':
        query = request.POST.get('keyword')
        # Filter ambulance items based on 'name' or 'manufacturer' containing the query
        ambulance_items = AmbulanceItem.objects.filter(Q(name__icontains=query) | Q(manufacturer__icontains=query))
        print('existe',ambulance_items)
        return render(request, 'Ambulance/AmbulanceItems/ambulance_Item.html', {'ambulance_items': ambulance_items})
    else:
        ambulance_items = AmbulanceItem.objects.all()
        return render(request, 'Ambulance/AmbulanceItems/ambulance_Item.html', {'ambulance_items': ambulance_items})
 #######################AmbulanceMedicine#############################
 
def ambulance_medicine_list(request):
    # Retrieve all medicine items from the database and pass them to the template
    ambulance_medicines = AmbulanceMedicine.objects.all().order_by('name')
    return render(request, 'Ambulance/AmbulanceMedicine/ambulance_medicines.html', {'ambulance_medicines': ambulance_medicines})


def create_ambulance_medicine(request): 
    if request.method == 'POST':
        form = AmbulanceMedicineForm(request.POST or None, request.FILES or None)  
        if form.is_valid():
            form.save()
            return redirect('ambulance_medicines')  
        else:
            print("Form errors:", form.errors)  # Debug print to check form errors
    else:
        form = AmbulanceMedicineForm()

    return render(request, 'Ambulance/AmbulanceMedicine/create_ambulance_medicine.html', {'form': form})

def update_ambulance_medicine(request, pk):
    ambulance_medicine = get_object_or_404(AmbulanceMedicine, pk=pk)
    if request.method == 'POST':
        form = AmbulanceMedicineForm(request.POST or None, request.FILES or None, instance=ambulance_medicine)
        if form.is_valid():
            form.save()
            return redirect('ambulance_medicines')
    else:
        form = AmbulanceMedicineForm(instance=ambulance_medicine)
    return render(request, 'Ambulance/AmbulanceMedicine/update_ambulance_medicine.html', {'form': form, 'ambulance_medicine': ambulance_medicine})

def view_ambulance_medicine(request, pk):
    ambulance_medicine = get_object_or_404(AmbulanceMedicine, pk=pk)
    form = AmbulanceMedicineForm(instance=ambulance_medicine)
    return render(request, 'Ambulance/AmbulanceMedicine/view_ambulance_medicine.html', {'form': form, 'ambulance_medicine': ambulance_medicine})

def delete_ambulance_medicine(request, pk):
    ambulance_medicine = get_object_or_404(AmbulanceMedicine, pk=pk)
    ambulance_medicine.delete()
    return redirect('ambulance_medicines')

def list_ambulance_medicines(request):
    # Retrieve all ambulance items from the database and pass them to the template
    ambulance_medicine = AmbulanceMedicine.objects.all().order_by('name')
    # Create a paginator with 10 items per page
    paginator = Paginator(ambulance_medicine, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    return render(request, 'list_ambulance_medicine.html', {"page_obj": page_obj})
#####################################

def search_AmbulanceMedicine(request):
    if request.method == 'POST':
        query = request.POST.get('keyword')

        ambulance_medicines = AmbulanceMedicine.objects.filter(Q(name__icontains=query) | Q(manufacturer__icontains=query))
        return render(request, 'Ambulance/AmbulanceMedicine/ambulance_medicines.html', {'ambulance_medicines': ambulance_medicines})
    else:
        ambulance_medicines = AmbulanceMedicine.objects.all()
        return render(request, 'Ambulance/AmbulanceMedicine/ambulance_medicines.html', {'ambulance_medicines': ambulance_medicines})   
#######################################################
from django.shortcuts import render, redirect
from .models import  VisiteMedicaleSpontanee
from .form import VisiteMedicaleSpontaneeForm

def visites_medicales_spontanees(request):
    visits = VisiteMedicaleSpontanee.objects.all().order_by('-visit_date')
    return render(request, 'visiteSpontanee/visits.html', {'visits': visits})


def create_spontaneous_visit(request):
    if request.method == 'POST':
        form = VisiteMedicaleSpontaneeForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('visits')  # Rediriger vers la liste des visites après l'enregistrement
        else:
            print("Form errors:", form.errors)  # Debug print to check form errors
    else:
        form = VisiteMedicaleSpontaneeForm()

    return render(request, 'visiteSpontanee/create_visit.html', {'form': form})

def view_visit(request, pk):
    visit = get_object_or_404(VisiteMedicaleSpontanee, pk=pk)
    form = VisiteMedicaleSpontaneeForm(instance=visit)
    return render(request, 'visiteSpontanee/view_visit.html', {'form': form, 'visit': visit})
###################################


def view_histoire_visite(request, pk):
    # Récupérer le patient à partir de la clé primaire (pk)
    patient = get_object_or_404(Patient, pk=pk)

    # Récupérer toutes les visites médicales associées à ce patient
    visits = VisiteMedicaleSpontanee.objects.filter(patient=patient)

    # Rendre la page 'view_histoire_visite.html' avec les visites médicales du patient
    return render(request, 'visiteSpontanee/visits.html', { 'visits': visits})
###################################################CheckIN##############################################################
from django.shortcuts import render, redirect, get_object_or_404
from .models import AmbulanceItem, AmbulanceMedicine, CheckInHistory
from .form import CheckInHistoryForm
def liste_controle(request):
  
    ambulance_items = AmbulanceItem.objects.all()
    ambulance_medicines = AmbulanceMedicine.objects.all()

    check_in_Historys = CheckInHistory.objects.all()

    return render(request, 'Ambulance/checkInAmbulanceItems.html', {'ambulance_items': ambulance_items, 'ambulance_medicines': ambulance_medicines, 'check_in_Historys': check_in_Historys})

from django.shortcuts import render, redirect
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
from .models import AmbulanceItem, AmbulanceMedicine, CheckInHistory
from .form import CheckInHistoryForm
from datetime import datetime
from django.templatetags.static import static 
from django.shortcuts import render, redirect
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image,KeepTogether
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
from .models import AmbulanceItem, AmbulanceMedicine, CheckInHistory
from .form import CheckInHistoryForm
from datetime import datetime

def enregistrer_liste_controle(request):
    if request.method == 'POST':
        # Generate and save the PDF file
        pdf_buffer = BytesIO()
        page_size = (500, 700)  # Adjust width and height as needed
        doc = SimpleDocTemplate(pdf_buffer, pagesize=page_size, leftMargin=5, rightMargin=5, topMargin=2, bottomMargin=2)
    
        elements = []
        company_logo_url = request.build_absolute_uri(static('nomac.jpg'))  # Build the absolute URL
        logo = Image(company_logo_url, width=100, height=70)  
        title_style = getSampleStyleSheet()["Title"]
        title1 = Paragraph("Ambulance Weekly ", title_style)
        title2= Paragraph("Inspection Checklist", title_style)
        current_date = datetime.now().strftime('%Y-%m-%d')
        date_style = getSampleStyleSheet()["Normal"]
        date = Paragraph(f"Date: {current_date}", date_style)
        logo_title_table = Table([[logo, title1]], colWidths=[5, '*'])
        header_content = [
            KeepTogether([logo_title_table]),
            Spacer(0, 5),
            KeepTogether([title2]),
            Spacer(0, 5),
            KeepTogether([date]),
        ]

        # Add the logo, title, and current date within a spacer to center align them
        elements.extend(header_content)

        # Get all AmbulanceItems and AmbulanceMedicine
        ambulance_items = AmbulanceItem.objects.all()
        ambulance_medicines = AmbulanceMedicine.objects.all()

        # Create a list of items and medicines in a table format
        items_header = ['SL', 'Items', 'Check In', 'Date', 'Remark']
        all_items = [items_header]
        remarks = []
        all_items_medicines=[]
        i = 0
        for item in ambulance_items:
            i += 1
            checked_status = "Existe" if request.POST.get(f"ambulance_item_{item.pk}") else "Not Existe"
            remark_item = request.POST.get(f"remark_item_{item.pk}")
            remarks.append(remark_item)
            all_items.append([i, item.name, checked_status, current_date, remark_item])

        # Apply styles for Items section
        table_style = TableStyle([
            ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
            ('BACKGROUND', (0, 0), (-1, 0), '#336699'),  # Header background color for Items
            ('TEXTCOLOR', (0, 0), (-1, 0), '#FFFFFF'),  # Header text color for Items
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('GRID', (0, 0), (-1, -1), 1, '#000000'),
        ])

        # Header for Medicine section
        medicine_header = ['SL', 'Medicine', 'Check In', 'Date', 'Remark']
        all_medicines=[medicine_header] # Add Medicine header row

        # Populate Medicine section
        i = 0
        for medicine in ambulance_medicines:
            i += 1 
            checked_status = "Existe" if request.POST.get(f"ambulance_medicine_{medicine.pk}") else "Not Existe"
            remark_medicine = request.POST.get(f"remark_medicine_{medicine.pk}")
            remarks.append(remark_medicine)
            all_medicines.append([i, medicine.name, checked_status, current_date, remark_medicine])
        all_items_medicines.append(all_items)
        all_items_medicines.append(all_medicines)
        # Apply styles for Medicine section
       
        column_widths = [30,200, 100, 80, 50]
        if all_items_medicines:
            items_table = Table(all_items, colWidths=column_widths)
            items_table.setStyle(table_style)
            elements.append(items_table)
            medicine_table = Table(all_medicines ,colWidths=column_widths)
            medicine_table.setStyle(table_style)
            elements.append(medicine_table)
        doc.build(elements)
            # Save the PDF to the CheckInHistory model
        form = CheckInHistoryForm(request.POST, request.FILES)
        if form.is_valid():
            check_in_history = form.save(commit=False)
            check_in_history.date_sauvegarde = datetime.now()
            check_in_history.file_pdf.save('checklist_report.pdf', pdf_buffer)
            check_in_history.remark = remarks
            check_in_history.save()
    # After saving, redirect to the history page
    return redirect('historique_sauvegardes')

def enregistrer_fichier(request, pk):
    # Get the CheckInHistory object using the primary key (pk)
    check_in_history = get_object_or_404(CheckInHistory, pk=pk)

    if request.method == 'POST' and request.FILES.get('other_file'):
        # Save the uploaded other file to the CheckInHistory entry
        uploaded_other_file = request.FILES['other_file']
        check_in_history.other_file = uploaded_other_file
        check_in_history.save()

    # Redirect back to the history page after the file is uploaded
    return redirect('historique_sauvegardes')

def historique_sauvegardes(request):
    # Récupérer tous les enregistrements de sauvegarde
    check_in_Historys = CheckInHistory.objects.all().order_by('-date_sauvegarde')
    return render(request, 'Ambulance/historique_sauvegardes.html', {'check_in_Historys': check_in_Historys})   

######################STOCKOUT############################""

from django.shortcuts import render, redirect
from .models import  OutStockMedicine
from .form import OutStockMedicineForm

def outStockMedicine(request):
    outMedicines = OutStockMedicine.objects.all()
    return render(request, 'Medicine/OUT/outMedicines.html', {'outMedicines': outMedicines})


def create_OutStockMedicine(request):
    if request.method == 'POST':
        form = OutStockMedicineForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('outMedicines')  
        else:
            print("Form errors:", form.errors)  
    else:
        form = OutStockMedicineForm()

    return render(request, 'Medicine/OUT/create_outMedicine.html', {'form': form})
def search_OutStockMedicine(request):
    if request.method == 'POST':
        query = request.POST.get('keyword')

        outMedicines = OutStockMedicine.objects.filter(Q(patient__name__icontains=query) | Q(medicine__name__icontains=query))
        return render(request, 'Medicine/OUT/outMedicines.html',{'outMedicines':outMedicines})
    else:
        outMedicines = OutStockMedicine.objects.all()
        return render(request, 'Medicine/OUT/outMedicines.html',{'outMedicines':outMedicines})
 
################################################################################################
#################################################################################################


def inde(request):
    notes = Notes.objects.all()
    return render(request, "Notes/index.html", {"notes": notes})


def new_note(request):
    form = NotesForm()
    if request.method == 'POST':
        form = NotesForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect("index")
    return render(request, "Notes/update.html", {"form": form})


def note_detail(request, pk):
    note = Notes.objects.get(id=pk)
    form = NotesForm(instance=note)
    if request.method == 'POST':
        form = NotesForm(request.POST, instance=note)
        if form.is_valid():
            form.save()
            return redirect("index")
    return render(request, "Notes/update.html", {"note": note, "form": form})


# def delete_note(request, pk):
def delete_note(request, pk):
    note = Notes.objects.get(id=pk)
    form = NotesForm(instance=note)
    if request.method == 'POST':
        note.delete()
        messages.info(request, "The note has been deleted")
        return redirect("index")
    return render(request, "Notes/delete.html", {"note": note, "form": form})


def search_page(request):
    if request.method == 'POST':
        search_text = request.POST['search']
        notes = Notes.objects.filter(heading__icontains=search_text) | Notes.objects.filter(text__icontains=search_text)

        return render(request, "Notes/search.html", {"notes": notes})

#########################dossier medical##################################
def dossier(request):
    return render(request, "dossier/dossier.html")


def patients_without_info_view(request):

    patients_without_info = Patient.objects.filter(info__isnull=True)
    print(patients_without_info)

    context = {
        'patients_without_info': patients_without_info
    }

    return render(request, 'dossier/Ajouterdossier.html', context)

def patients_with_info_view(request):

    patients_with_info = Patient.objects.filter(info__isnull=False)
    context = {
        'patients_with_info': patients_with_info
    }

    return render(request, 'dossier/patient_dossier.html', context)

def add_dossier(request):
    return render(request, 'Medecin/medecin.html')



def add_info(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)

    if request.method == 'POST':
        form = InfoForm(request.POST)
        if form.is_valid():
            info = form.save(commit=False)
            info.patient = patient
            info.save()
            return redirect('ajouterdossier')
    else:
        form = InfoForm()
    context = {
        'patient': patient,
        'form': form,
    }
    return render(request, 'dossier/menu_content1.html', context)

def menu1_view(request):
    if request.method == 'POST':
        form = InfoForm(request.POST)
        if form.is_valid():
            print("Form data:", form.cleaned_data)
            form.save()
            return redirect('menu1')
        else:
            print("Form errors:", form.errors)
    else:
        form = InfoForm()
    return render(request, 'dossier1/menu_content1.html',{'form':form})



def menu3_view(request):
    if request.method == 'POST':
        form = SystematiqueForm(request.POST)
        if form.is_valid():
            print("Form data:", form.cleaned_data)
            form.save()
            return redirect('add_dossier')
        else:
            print("Form errors:", form.errors)
    else:
        form = SystematiqueForm()
    return render(request, 'dossier/VMR.html', {'form': form})


from django.shortcuts import render, get_object_or_404, redirect
from .models import Patient, VisiteReprise

################################### VMR #################################################################



def visite_reprise(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)
    locale.setlocale(locale.LC_TIME, 'fr_FR.utf8')
    if request.method == 'POST':
        form = VisiteRepriseForm(request.POST)
        if form.is_valid():
            visite = form.save(commit=False)
            visite.patient = patient
            visite.save()
            return redirect('list_reprise', patient_id=patient_id)
            messages.success(request, "La visite de reprise a été enregistrée avec succès.")
        else:
            messages.error(request, "Veuillez remplir tous les champs requis.")
    else:
        form = VisiteRepriseForm(initial={'patient': patient})

    return render(request, 'dossier1/VMR/ajouter_vmr.html', {'form': form, 'patient': patient})


def reprise_list(request):
    patients = Patient.objects.all()
    return render(request, 'dossier1/VMR/VMR.html', {'patients': patients})

def liste_visites_par_patient(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)
    visites = patient.visites_reprise.all()

    return render(request, 'dossier1/VMR/liste_visite.html', {'patient': patient, 'visites': visites})


def detail_visite_reprise(request, patient_id, date_visite):
    patient = get_object_or_404(Patient, pk=patient_id)
    visite = get_object_or_404(VisiteReprise, patient=patient, date=date_visite)
    visites_patient_date = VisiteReprise.objects.filter(patient=patient, date=date_visite)

    return render(request, 'dossier/VMR/voir_visite.html',
                  {'visite': visite, 'visites_patient_date': visites_patient_date})

######################analyse#############

def liste_analyses(request):
    analyses = Analyse.objects.all()
    data = [{'id': analyse.id, 'nom': analyse.nom} for analyse in analyses]
    return JsonResponse(data, safe=False)
###############################VMS########################



def ajouter_vms(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)

    if request.method == 'POST':
        form = SystematiqueForm(request.POST)
        if form.is_valid():
            new_date = form.cleaned_data['date']
            if systematique.objects.filter(patient=patient, date=new_date).exists():
                form.add_error('date', 'une visite dans la meme date existe déja.')
            else:
                new_systematique = form.save(commit=False)
                new_systematique.patient = patient
                new_systematique.save()


                selected_analyses = form.cleaned_data['analyses']
                new_systematique.analyses.set(selected_analyses)

                return redirect('liste_visite_par_patient', patient_id=patient_id)
        else:
            print("Form errors:", form.errors)
    else:
        form = SystematiqueForm(initial={'patient': patient})

    return render(request, 'dossier/VMS/ajouter_VMS.html', {'form': form, 'patient': patient})

def sys_list(request):
    patients = Patient.objects.all()
    return render(request, 'dossier/VMS/VMS.html', {'patients': patients})

def liste_visite_par_patient(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)
    visites = patient.systematique_set.all()  # Utilisez la relation "systematique_set" pour récupérer les visites
    url = reverse('liste_visite_par_patient', args=[patient.id])
    return render(request, 'dossier/VMS/liste_vms.html', {'patient': patient, 'visites': visites})

def detail_vms(request, patient_id, date_visite):
    patient = get_object_or_404(Patient, pk=patient_id)
    visite = get_object_or_404(systematique, patient=patient, date=date_visite)
    visites_patient_date = systematique.objects.filter(patient=patient, date=date_visite)

    return render(request, 'dossier/VMS/voir_vms.html',
                  {'visite': visite, 'visites_patient_date': visites_patient_date})


def delete_vms(request, patient_id, visite_id):
    visite = get_object_or_404(systematique, pk=visite_id, patient_id=patient_id)
    if request.method == 'POST':
        visite.delete()
        return redirect('liste_visite_par_patient', patient_id=patient_id)
    return render(request, 'dossier/VMS/supprimer_visite.html', {'visite': visite})



##################################dossier##################################
def view_info(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    info_instance = get_object_or_404(info, patient=patient)
    return render(request, 'dossier/view_info.html', {'patient': patient, 'info': info_instance})




def view_patient_dossier(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    form = PatientForm(instance=patient)
    print("Patient PK:", patient.pk)
    return render(request, 'dossier/patient.html', {'form': form , 'patient':patient})

def liste_vms_par_patient(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    visites = patient.systematique_set.all()  # Utilisez la relation "systematique_set" pour récupérer les visites
    url = reverse('liste_visite_par_patient', args=[patient.id])

    return render(request, 'dossier/list_vms.html', {'patient': patient, 'visites': visites})

def detail_of_vms(request, patient_id, date_visite):
    patient = get_object_or_404(Patient, pk=patient_id)
    visite = get_object_or_404(systematique, patient=patient, date=date_visite)
    visites_patient_date = systematique.objects.filter(patient=patient, date=date_visite)
    return render(request, 'dossier/vms.html',
                  {'visite': visite, 'visites_patient_date': visites_patient_date})



def detail_vmr(request, patient_id, date_visite):
    patient = get_object_or_404(Patient, pk=patient_id)
    visite = get_object_or_404(VisiteReprise, patient=patient, date=date_visite)


    visites_patient_date = VisiteReprise.objects.filter(patient=patient, date=date_visite)

    return render(request, 'dossier/vmr.html',
                  {'visite': visite, 'visites_patient_date': visites_patient_date})

def liste_vmr_par_patient(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    visites = patient.visites_reprise.all()

    return render(request, 'dossier/list_vmr.html', {'patient': patient, 'visites': visites})


def liste_surv_par_patient(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    visites = patient.surveillance_set.all()  # Utilisez la relation "systematique_set" pour récupérer les visites
    url = reverse('liste_visite_par_patient', args=[patient.id])

    return render(request, 'dossier/list_surveillance.html', {'patient': patient, 'visites': visites})

def detail_surv(request, patient_id, date_visite):
    patient = get_object_or_404(Patient, pk=patient_id)
    visite = get_object_or_404(surveillance, patient=patient, date=date_visite)

    visites_patient_date = surveillance.objects.filter(patient=patient, date=date_visite)

    return render(request, 'dossier/surv.html',
                  {'visite': visite, 'visites_patient_date': visites_patient_date})

##################labo#################""""

def patients_sans_laboratoire(request):
    visites_sans_laboratoire = systematique.objects.filter(laboratoire__isnull=True)
    return render(request, 'dossier/Laboratoire/list_labo.html', {'visites_sans_laboratoire': visites_sans_laboratoire})

def ajouter_resultat_laboratoire(request, visite_id, analyse_id):
    laboratoire = Laboratoire.objects.filter(systematique_id=visite_id, analyse_id=analyse_id).first()
    if laboratoire:
        fichier_path = laboratoire.fichier.path
        response = FileResponse(open(fichier_path, 'rb'))
        return response
    return HttpResponse("Aucun fichier de laboratoire trouvé.")

def scanner_view(request, visite_id, analyse_id):
    visite = get_object_or_404(systematique, pk=visite_id)
    analyse = get_object_or_404(Analyse, pk=analyse_id)
    try:
        laboratoire = Laboratoire.objects.get(systematique=visite, analyse=analyse)
    except Laboratoire.DoesNotExist:
        laboratoire = None

    context = {
        'visite': visite,
        'analyse': analyse,
        'laboratoire': laboratoire
    }

    return render(request, 'dossier/Laboratoire/list_labo.html', context)


def upload_laboratoire(request, visite_id):
    visite = systematique.objects.get(pk=visite_id)
    analyses_with_files = []

    if request.method == 'POST':
        all_analyses_uploaded = True

        for analyse in visite.analyses.all():
            fichier_name = f'fichier_{analyse.pk}'
            if fichier_name in request.FILES:
                fichier = request.FILES[fichier_name]
                analyses_with_files.append(analyse)
            else:
                all_analyses_uploaded = False

        if all_analyses_uploaded and len(analyses_with_files) == visite.analyses.count():
            for analyse in analyses_with_files:
                fichier_name = f'fichier_{analyse.pk}'
                fichier = request.FILES[fichier_name]
                Laboratoire.objects.create(systematique=visite, analyse=analyse, fichier=fichier)
            return redirect('fiche_médicale')  # Replace with the actual URL
        else:
            error_message = "Merci de scanner tous les fichiers"
            context = {
                'visite': visite,
                'error_message': error_message,
            }
            return render(request, 'dossier/Laboratoire/ajouter_scan.html', context)


    context = {
        'visite': visite,
    }


    return render(request, 'dossier/Laboratoire/ajouter_scan.html', context)

#####################################surveillance###############################################


def add_surveillance(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)

    if request.method == 'POST':
        form = SurveillanceForm(request.POST)
        if form.is_valid():
            new_date = form.cleaned_data['date']
            if surveillance.objects.filter(patient=patient, date=new_date).exists():
                form.add_error('date', 'une visite dans la meme date existe déja.')
            else:
                new_surveillance = form.save(commit=False)
                new_surveillance.patient = patient
                new_surveillance.save()
                return redirect('liste_surveillance_par_patient', patient_id=patient_id)

    else:
        form = SurveillanceForm(initial={'patient': patient})

    return render(request,  'dossier/surveillance/add_surveillance.html', {'form': form, 'patient': patient})


def surveillance_list(request):
    patients = Patient.objects.all()
    return render(request, 'dossier/surveillance/surveillance.html', {'patients': patients})

def liste_surveillance_par_patient(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)
    visites = patient.surveillance_set.all()  # Utilisez la relation "systematique_set" pour récupérer les visites
    url = reverse('liste_visite_par_patient', args=[patient.id])

    return render(request, 'dossier/surveillance/list_surveillance.html', {'patient': patient, 'visites': visites})

def detail_surveillance(request, patient_id, date_visite):
    patient = get_object_or_404(Patient, pk=patient_id)
    visite = get_object_or_404(surveillance, patient=patient, date=date_visite)

    visites_patient_date = surveillance.objects.filter(patient=patient, date=date_visite)

    return render(request, 'dossier/surveillance/voir_surv.html',
                  {'visite': visite, 'visites_patient_date': visites_patient_date})
